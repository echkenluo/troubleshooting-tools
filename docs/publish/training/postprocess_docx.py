#!/usr/bin/env python3
"""Post-process pandoc docx: style code blocks with background, borders, monospace font."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

docx_path = sys.argv[1]
doc = Document(docx_path)

CODE_BG = 'F4F4F4'
CODE_BORDER_COLOR = 'D0D0D0'
CODE_FONT = 'Courier New'
CODE_FONT_SIZE = Pt(8.5)
INLINE_CODE_COLOR = RGBColor(0xd6, 0x33, 0x84)

def add_shading_to_paragraph(para, color_hex):
    """Add background shading to a paragraph element."""
    pPr = para._element.get_or_add_pPr()
    # Remove existing shading
    for existing in pPr.findall(qn('w:shd')):
        pPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)

def add_border_to_paragraph(para, color_hex, sides=None):
    """Add borders to a paragraph."""
    if sides is None:
        sides = ['left']
    pPr = para._element.get_or_add_pPr()
    # Remove existing borders
    for existing in pPr.findall(qn('w:pBdr')):
        pPr.remove(existing)
    pBdr = OxmlElement('w:pBdr')
    for side in sides:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8' if side == 'left' else '4')
        border.set(qn('w:space'), '6' if side == 'left' else '4')
        border.set(qn('w:color'), color_hex)
        pBdr.append(border)
    pPr.append(pBdr)

def set_paragraph_spacing(para, before=Pt(0), after=Pt(0), line_spacing=Pt(13)):
    """Set compact spacing for code paragraphs."""
    pf = para.paragraph_format
    pf.space_before = before
    pf.space_after = after
    pf.line_spacing = line_spacing

code_block_count = 0
inline_code_count = 0

# Process all paragraphs
for para in doc.paragraphs:
    style_name = para.style.name

    if style_name == 'Source Code':
        code_block_count += 1
        # Background shading
        add_shading_to_paragraph(para, CODE_BG)
        # Left accent border
        add_border_to_paragraph(para, '4A90D9', sides=['left'])
        # Compact spacing
        set_paragraph_spacing(para)
        # Indentation
        para.paragraph_format.left_indent = Inches(0.15)
        para.paragraph_format.right_indent = Inches(0.15)

        # Force monospace font on all runs
        for run in para.runs:
            run.font.name = CODE_FONT
            run.font.size = CODE_FONT_SIZE
            # Set East Asian font too (for CJK characters in comments)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:ascii'), CODE_FONT)
            rFonts.set(qn('w:hAnsi'), CODE_FONT)
            rFonts.set(qn('w:cs'), CODE_FONT)

    else:
        # Check for inline code runs
        for run in para.runs:
            if run.style and run.style.name == 'Verbatim Char':
                inline_code_count += 1
                run.font.name = CODE_FONT
                run.font.size = Pt(9)
                run.font.color.rgb = INLINE_CODE_COLOR
                # Add background via run shading
                rPr = run._element.get_or_add_rPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F0F0F0')
                rPr.append(shd)

# Also style tables for better appearance
for table in doc.tables:
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        # Add table borders
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)
        for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'BBBBBB')
            # Remove existing
            existing = tblBorders.find(qn(f'w:{side}'))
            if existing is not None:
                tblBorders.remove(existing)
            tblBorders.append(border)

    # Style header row (first row)
    if table.rows:
        for cell in table.rows[0].cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'E8E8E8')
            # Remove existing
            for existing in tcPr.findall(qn('w:shd')):
                tcPr.remove(existing)
            tcPr.append(shd)
            # Bold header text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

doc.save(docx_path)
print(f"Done: {code_block_count} code blocks styled, {inline_code_count} inline code runs styled")
